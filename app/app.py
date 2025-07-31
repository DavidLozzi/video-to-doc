import streamlit as st
import os
import tempfile
import cv2
import subprocess
import re
import numpy as np
from docx import Document
from docx.shared import Inches
import httpx
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import time

os.environ["OPENCV_FFMPEG_READ_ATTEMPTS"] = "8192"
if not os.path.exists("output"):
    os.makedirs("output")
print(os.getcwd())


def get_chapters_from_url(url):
    """Get all chapters from a landing page"""
    print(f"Getting chapters from URL: {url}")
    
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        with httpx.Client(timeout=30.0) as client:
            response = client.get(url, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Look for chapter links - more specific patterns for comic/manga sites
            chapters = []
            
            # Common navigation elements to skip
            skip_keywords = [
                'login', 'register', 'search', 'home', 'about', 'contact', 'privacy', 'terms',
                'comic', 'marvel', 'dc', 'manga', 'anime', 'genre', 'category', 'tag',
                'author', 'artist', 'publisher', 'status', 'rating', 'language',
                'menu', 'nav', 'header', 'footer', 'sidebar', 'breadcrumb',
                'facebook', 'twitter', 'instagram', 'youtube', 'discord',
                'donate', 'support', 'advertise', 'submit', 'upload'
            ]
            
            # Debug: Show all links found
            all_links = soup.find_all('a', href=True)
            print(f"Found {len(all_links)} total links on the page")
            
            # Pattern 1: Look for chapter/episode links with specific patterns
            chapter_links = soup.find_all('a', href=True)
            for link in chapter_links:
                href = link.get('href', '').lower()
                text = link.get_text().strip()
                text_lower = text.lower()
                
                # Skip navigation and non-chapter links
                if any(skip in href or skip in text_lower for skip in skip_keywords):
                    continue
                
                # Look for chapter indicators
                chapter_indicators = ['chapter', 'ch', 'episode', 'ep', 'issue', 'vol', 'volume']
                has_chapter_indicator = any(indicator in href or indicator in text_lower for indicator in chapter_indicators)
                
                # Look for numbered patterns (Chapter 1, Episode 2, etc.)
                has_number = re.search(r'\d+', text) or re.search(r'\d+', href)
                
                # Look for specific comic/manga patterns
                is_comic_pattern = (
                    re.search(r'chapter-\d+', href) or
                    re.search(r'episode-\d+', href) or
                    re.search(r'issue-\d+', href) or
                    re.search(r'vol-\d+', href) or
                    re.search(r'volume-\d+', href) or
                    (has_number and has_chapter_indicator)
                )
                
                if is_comic_pattern or (has_number and has_chapter_indicator):
                    # Convert relative URLs to absolute
                    if not href.startswith(('http://', 'https://')):
                        href = urljoin(url, href)
                    
                    # Clean up the title
                    if not text:
                        text = f"Chapter {len(chapters) + 1}"
                    
                    chapters.append({
                        'title': text,
                        'url': href
                    })
                    print(f"Found chapter: {text} -> {href}")
            
            # Pattern 2: Look for links in specific containers that might contain chapters
            if not chapters:
                print("No chapters found with pattern 1, trying pattern 2...")
                # Look for common chapter list containers
                chapter_containers = soup.find_all(['div', 'ul', 'ol'], class_=re.compile(r'chapter|episode|list', re.I))
                print(f"Found {len(chapter_containers)} potential chapter containers")
                
                for container in chapter_containers:
                    container_links = container.find_all('a', href=True)
                    for link in container_links:
                        href = link.get('href', '').lower()
                        text = link.get_text().strip()
                        
                        # Skip if it's clearly not a chapter
                        if any(skip in href or skip in text.lower() for skip in skip_keywords):
                            continue
                        
                        # Must have a number to be considered a chapter
                        if re.search(r'\d+', text) or re.search(r'\d+', href):
                            if not href.startswith(('http://', 'https://')):
                                href = urljoin(url, href)
                            
                            chapters.append({
                                'title': text or f"Chapter {len(chapters) + 1}",
                                'url': href
                            })
                            print(f"Found chapter in container: {text} -> {href}")
            
            # Pattern 3: Look for any numbered links that might be chapters (very strict)
            if not chapters:
                print("No chapters found with pattern 2, trying pattern 3...")
                for link in chapter_links:
                    href = link.get('href', '')
                    text = link.get_text().strip()
                    
                    # Must have a number and not be navigation
                    if (re.search(r'\d+', text) or re.search(r'\d+', href)) and not any(skip in href.lower() or skip in text.lower() for skip in skip_keywords):
                        # Additional check: URL should look like a chapter URL
                        if any(pattern in href for pattern in ['/chapter/', '/episode/', '/issue/', '/vol/', '/volume/']):
                            if not href.startswith(('http://', 'https://')):
                                href = urljoin(url, href)
                            
                            chapters.append({
                                'title': text or f"Chapter {len(chapters) + 1}",
                                'url': href
                            })
                            print(f"Found chapter with URL pattern: {text} -> {href}")
            
            # Pattern 4: More lenient - any link with a number that's not navigation
            if not chapters:
                print("No chapters found with pattern 3, trying pattern 4 (lenient)...")
                for link in chapter_links:
                    href = link.get('href', '')
                    text = link.get_text().strip()
                    
                    # Skip obvious navigation
                    if any(skip in href.lower() or skip in text.lower() for skip in ['login', 'register', 'search', 'home', 'about', 'contact', 'menu', 'nav']):
                        continue
                    
                    # Must have a number
                    if re.search(r'\d+', text) or re.search(r'\d+', href):
                        if not href.startswith(('http://', 'https://')):
                            href = urljoin(url, href)
                        
                        chapters.append({
                            'title': text or f"Chapter {len(chapters) + 1}",
                            'url': href
                        })
                        print(f"Found potential chapter (lenient): {text} -> {href}")
            
            # Pattern 5: Look for specific "Ch.X" patterns FIRST (like Ch.6, Ch.5, etc.)
            if not chapters:
                print("No chapters found with pattern 4, trying pattern 5 (Ch.X patterns)...")
                
                # Look for text containing "Ch." followed by a number
                ch_pattern = re.compile(r'Ch\.\s*\d+', re.IGNORECASE)
                
                # Search in all text elements and collect ALL matches
                all_elements = soup.find_all(text=True)
                found_chapters = set()  # Use set to avoid duplicates
                
                for element in all_elements:
                    text = element.strip()
                    if text and ch_pattern.search(text):
                        # Find ALL matches in this text
                        matches = ch_pattern.findall(text)
                        for match in matches:
                            chapter_text = match
                            if chapter_text not in found_chapters:
                                found_chapters.add(chapter_text)
                                chapters.append({
                                    'title': chapter_text,
                                    'url': url,  # Same URL since it's JS-based
                                })
                                print(f"Found Ch.X pattern: {chapter_text}")
                
                # Also look for elements containing "Ch." patterns
                ch_elements = soup.find_all(text=ch_pattern)
                for element in ch_elements:
                    text = element.strip()
                    if text:
                        matches = ch_pattern.findall(text)
                        for match in matches:
                            chapter_text = match
                            if chapter_text not in found_chapters:
                                found_chapters.add(chapter_text)
                                chapters.append({
                                    'title': chapter_text,
                                    'url': url,
                                })
                                print(f"Found Ch.X element: {chapter_text}")
            
            # Pattern 6: Look for other JavaScript-based chapter selectors (for SPAs)
            if not chapters:
                print("No chapters found with pattern 5, trying pattern 6 (other JS-based)...")
                
                # Look for elements that might be chapter selectors
                potential_chapters = []
                
                # Look for elements with data attributes that might indicate chapters
                data_elements = soup.find_all(attrs={"data-": True})
                for element in data_elements:
                    text = element.get_text().strip()
                    if text and re.search(r'\d+', text):
                        # Skip if it's clearly not a chapter (like comic titles with years)
                        if not any(skip in text.lower() for skip in ['star wars', 'marvel', 'comic', '2025', '2024', '2023']):
                            potential_chapters.append({
                                'title': text,
                                'url': url,  # Same URL since it's JS-based
                                'element': element
                            })
                
                # Look for elements with specific classes that might be chapters
                chapter_classes = ['chapter', 'episode', 'issue', 'volume', 'part']
                for class_name in chapter_classes:
                    elements = soup.find_all(class_=re.compile(class_name, re.I))
                    for element in elements:
                        text = element.get_text().strip()
                        if text and re.search(r'\d+', text):
                            potential_chapters.append({
                                'title': text,
                                'url': url,  # Same URL since it's JS-based
                                'element': element
                            })
                
                # Look for numbered buttons or divs
                numbered_elements = soup.find_all(['button', 'div', 'span'], text=re.compile(r'\d+'))
                for element in numbered_elements:
                    text = element.get_text().strip()
                    if text and re.search(r'\d+', text):
                        # Skip if it's clearly navigation or comic title
                        if not any(skip in text.lower() for skip in ['login', 'register', 'search', 'home', 'about', 'contact', 'menu', 'nav', 'star wars', 'marvel', 'comic', '2025', '2024', '2023']):
                            potential_chapters.append({
                                'title': text,
                                'url': url,  # Same URL since it's JS-based
                                'element': element
                            })
                
                # Remove duplicates and add to chapters
                seen_titles = set()
                for item in potential_chapters:
                    if item['title'] not in seen_titles:
                        seen_titles.add(item['title'])
                        chapters.append({
                            'title': item['title'],
                            'url': item['url']
                        })
                        print(f"Found JS-based chapter: {item['title']} -> {item['url']}")
            
            # Remove duplicates based on URL
            seen_urls = set()
            unique_chapters = []
            for chapter in chapters:
                if chapter['url'] not in seen_urls:
                    seen_urls.add(chapter['url'])
                    unique_chapters.append(chapter)
            
            # Sort chapters by number if possible
            def extract_number(title):
                numbers = re.findall(r'\d+', title)
                return int(numbers[0]) if numbers else 0
            
            unique_chapters.sort(key=lambda x: extract_number(x['title']))
            
            print(f"Final result: Found {len(unique_chapters)} unique chapters")
            for chapter in unique_chapters:
                print(f"  - {chapter['title']}: {chapter['url']}")
            
            return unique_chapters
        
    except Exception as e:
        st.error(f"Failed to get chapters from URL: {e}")
        print(f"Error getting chapters: {e}")
        return []


def download_images_from_url(url):
    """Download all images from a web page"""
    print(f"Downloading images from URL: {url}")
    
    try:
        # Basic headers for initial request
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
        }
        
        with httpx.Client(timeout=30.0) as client:
            # First, get the main page to extract dynamic headers and comic ID
            response = client.get(url, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            html_content = response.text
            
            # Extract dynamic Next.js headers from the HTML
            next_action = None
            next_router_state_tree = None
            
            # Look for next-action in script tags or JSON data
            import re
            na_match = re.search(r'"nextAction"\s*:\s*"([^"]+)"', html_content)
            if na_match:
                next_action = na_match.group(1)
                print(f"Found next-action: {next_action}")
            
            # Look for router state tree
            nrst_match = re.search(r'"tree"\s*:\s*(\[[^\]]+\])', html_content)
            if nrst_match:
                next_router_state_tree = nrst_match.group(1)
                print(f"Found router state tree: {next_router_state_tree[:100]}...")
            
            # Also try to find them in __NEXT_DATA__ script tag
            next_data_script = soup.find('script', {'id': '__NEXT_DATA__'})
            if next_data_script and next_data_script.string:
                try:
                    import json
                    next_data = json.loads(next_data_script.string)
                    if 'nextAction' in next_data:
                        next_action = next_data['nextAction']
                        print(f"Found next-action in __NEXT_DATA__: {next_action}")
                    if 'tree' in next_data:
                        next_router_state_tree = json.dumps(next_data['tree'])
                        print(f"Found router state tree in __NEXT_DATA__")
                except json.JSONDecodeError:
                    pass
            
            # Extract comic ID from the page (look for patterns like [2779508])
            comic_id = None
            script_tags = soup.find_all('script')
            for script in script_tags:
                if script.string:
                    id_pattern = r'\[(\d+)\]'
                    matches = re.findall(id_pattern, script.string)
                    if matches:
                        comic_id = matches[0]
                        print(f"Found comic ID: {comic_id}")
                        break
            
            image_urls = []
            
            # If we found the required headers and comic ID, make the server action request
            if comic_id and next_action and next_router_state_tree:
                print(f"Making server action request with comic ID: {comic_id}")
                action_data = f'[{comic_id}]'
                
                # Set up headers for the server action request
                action_headers = {
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36',
                    'accept': 'text/x-component',
                    'content-type': 'text/plain;charset=UTF-8',
                    'next-action': next_action,
                    'next-router-state-tree': next_router_state_tree,
                    'referer': url,
                    'origin': 'https://freecomic.to'
                }
                
                try:
                    action_response = client.post(url, headers=action_headers, data=action_data)
                    action_response.raise_for_status()
                    print(f"Server action response: {action_response.text[:500]}...")
                    
                    # Parse the JSON response to extract image URLs
                    try:
                        # The response format is: "0:{"a":"$@1","f":"","b":"paBJs_QLFv962VoaR_i-s"}\n1:[{...}]"
                        response_lines = action_response.text.strip().split('\n')
                        if len(response_lines) >= 2:
                            # Get the second line which contains the image array
                            image_data_line = response_lines[1]
                            if image_data_line.startswith('1:'):
                                # Extract the JSON array part
                                json_part = image_data_line[2:]  # Remove "1:" prefix
                                image_array = json.loads(json_part)
                                
                                # Extract src URLs from each image object
                                for image_obj in image_array:
                                    if 'src' in image_obj:
                                        image_urls.append(image_obj['src'])
                                        print(f"Found image URL: {image_obj['src']}")
                                
                                print(f"Found {len(image_array)} images in server action response")
                            else:
                                print("Unexpected response format")
                    except json.JSONDecodeError as e:
                        print(f"Failed to parse JSON response: {e}")
                        # Fallback to regex pattern matching
                        imgor_pattern = r'https://imgor\.comick\.site/[^"\s]+'
                        matches = re.findall(imgor_pattern, action_response.text)
                        image_urls.extend(matches)
                        print(f"Found {len(matches)} imgor URLs using regex fallback")
                    
                except Exception as e:
                    print(f"Server action request failed: {e}")
            else:
                print(f"Missing required data: comic_id={comic_id}, next_action={next_action}, router_tree={next_router_state_tree is not None}")
            
            # Also check the main page for images as fallback
            img_tags = soup.find_all('img')
            
            # Look for imgor.comick.site URLs in script tags
            for script in script_tags:
                if script.string:
                    imgor_pattern = r'https://imgor\.comick\.site/[^"\s]+'
                    matches = re.findall(imgor_pattern, script.string)
                    image_urls.extend(matches)
                    print(f"Found {len(matches)} imgor URLs in script")
            
            # Look for any data attributes that might contain image URLs
            for img in img_tags:
                # Check all possible image attributes
                for attr in ['src', 'data-src', 'data-original', 'data-lazy', 'data-image']:
                    value = img.get(attr)
                    if value and 'imgor.comick.site' in value:
                        image_urls.append(value)
                        print(f"Found imgor URL in {attr}: {value}")
            
            # Add regular image sources
            for img in img_tags:
                src = img.get('src')
                data_src = img.get('data-src')
                
                if src:
                    image_urls.append(src)
                if data_src:
                    image_urls.append(data_src)
            
            # Remove duplicates
            image_urls = list(set(image_urls))
            
            if not image_urls:
                st.error("No images found on the page")
                return []
            
            # Create temporary directory for downloaded images
            temp_dir = tempfile.mkdtemp(prefix="web_images_")
            downloaded_images = []
            
            st.text(f"Found {len(image_urls)} images, downloading...")
            progress_bar = st.progress(0)
            
            # Debug: Print all image sources found
            print("All image sources found:")
            for i, img_url in enumerate(image_urls):
                print(f"  Image {i}: {img_url}")
            
            for i, img_url in enumerate(image_urls):
                if not img_url:
                    continue
                    
                # Convert relative URLs to absolute
                if not img_url.startswith(('http://', 'https://')):
                    img_url = urljoin(url, img_url)
                
                print(f"Downloading image {i}: {img_url}")
                
                try:
                    img_response = client.get(img_url, headers=headers, timeout=10.0)
                    img_response.raise_for_status()
                    
                    # Determine file extension from content type or URL
                    content_type = img_response.headers.get('content-type', '')
                    if 'jpeg' in content_type or 'jpg' in content_type:
                        ext = '.jpg'
                    elif 'png' in content_type:
                        ext = '.png'
                    elif 'gif' in content_type:
                        ext = '.gif'
                    elif 'webp' in content_type:
                        ext = '.webp'
                    else:
                        # Try to get extension from URL
                        parsed_url = urlparse(img_url)
                        path = parsed_url.path
                        if '.' in path:
                            ext = '.' + path.split('.')[-1].lower()
                        else:
                            ext = '.jpg'  # Default to jpg
                    
                    # Save image
                    img_filename = f"web_image_{i:04d}{ext}"
                    img_path = os.path.join(temp_dir, img_filename)
                    
                    with open(img_path, 'wb') as f:
                        f.write(img_response.content)
                    
                    downloaded_images.append(img_path)
                    
                    # Update progress
                    progress = (i + 1) / len(image_urls)
                    progress_bar.progress(progress)
                    
                except Exception as e:
                    print(f"Failed to download image {img_url}: {e}")
                    continue
            
            progress_bar.empty()
            st.text(f"Successfully downloaded {len(downloaded_images)} images")
            
            return downloaded_images, temp_dir
        
    except Exception as e:
        st.error(f"Failed to download images from URL: {e}")
        return [], None


def is_frame_unique_pixel(current_frame, previous_unique_frame, frame_index):
    """Check if frame is unique using direct pixel comparison"""
    if previous_unique_frame is None:
        return True

    # Convert both frames to grayscale and resize for comparison
    current_gray = cv2.cvtColor(current_frame, cv2.COLOR_BGR2GRAY)
    previous_gray = cv2.cvtColor(previous_unique_frame, cv2.COLOR_BGR2GRAY)
    
    # Resize both to same size for comparison
    current_resized = cv2.resize(current_gray, (100, 100))
    previous_resized = cv2.resize(previous_gray, (100, 100))
    
    # Calculate mean absolute difference
    diff = np.mean(np.abs(current_resized.astype(np.float32) - previous_resized.astype(np.float32)))
    
    # Use HIGH threshold - only frames that are very different are considered unique
    is_unique = diff > 70.0  # Much higher threshold - less sensitive
    
    print(f"Frame {frame_index}: Pixel diff = {diff:.2f} (unique: {is_unique})")
    
    return is_unique


def process_images_from_url(url):
    """Process images downloaded from URL like a video"""
    print("Processing images from URL...")
    
    # Download images from URL
    downloaded_images, temp_dir = download_images_from_url(url)
    
    if not downloaded_images:
        return []
    
    try:
        st.text("Processing all images")
        progress_bar = st.progress(0)
        status = st.empty()
        
        output_images = []
        
        cols = [st.columns(3) for _ in range(len(downloaded_images) // 3 + 1)]
        col_index = 0
        row_index = 0
        
        for img_index, img_path in enumerate(downloaded_images):
            percentage_done = (img_index / len(downloaded_images)) * 100
            progress_bar.progress(int(percentage_done))
            
            # Load the image
            frame = cv2.imread(img_path)
            if frame is None:
                continue
            
            # Save to output directory (no uniqueness check)
            output_img_path = f"./output/web_image_{img_index:04d}.jpg"
            cv2.imwrite(output_img_path, frame)
            output_images.append(output_img_path)
            status.text(f"Processed image {img_index}")
            
            # Display in Streamlit
            if row_index < len(cols) and col_index < 3:
                cols[row_index][col_index].image(output_img_path, width=300)
                col_index += 1
                if col_index == 3:
                    col_index = 0
                    row_index += 1
        
        # Clean up temporary directory
        if temp_dir:
            for file in os.listdir(temp_dir):
                try:
                    os.unlink(os.path.join(temp_dir, file))
                except:
                    pass
            try:
                os.rmdir(temp_dir)
            except:
                pass
        
        return output_images
        
    except Exception as e:
        print(f"Error during image processing: {e}")
        return []


def extract_unique_frames(uploaded_file):
    print("Extracting unique frames from video...")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp:
        file_bytes = uploaded_file.read()
        temp.write(file_bytes)
        temp_file_name = temp.name

    # Create temporary directories for frame extraction
    temp_all_folder = tempfile.mkdtemp(prefix="frames_all_")
    temp_unique_folder = tempfile.mkdtemp(prefix="frames_unique_")

    try:
        # Use FFmpeg to extract frames at intervals (every 6th frame for more samples)
        ffmpeg_cmd = [
            "ffmpeg",
            "-i",
            temp_file_name,
            "-vf",
            "select=not(mod(n\\,6))",
            "-vsync",
            "vfr",
            "-frame_pts",
            "true",
            f"{temp_all_folder}/frame_%05d.jpg",
        ]

        try:
            subprocess.run(ffmpeg_cmd, capture_output=True, text=True, check=True)
            print("Frame extraction completed successfully")
        except FileNotFoundError:
            st.error("FFmpeg not found. Please install ffmpeg first.")
            return []
        except subprocess.CalledProcessError as e:
            st.error(f"FFmpeg failed: {e.stderr}")
            return []

        # Get video stats
        cap = cv2.VideoCapture(temp_file_name)
        fps = cap.get(cv2.CAP_PROP_FPS)
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        cap.release()

        # Get all frame files sorted by number
        frame_files = sorted(
            [f for f in os.listdir(temp_all_folder) if f.startswith("frame_")]
        )

        st.text("Finding unique frames")
        progress_bar = st.progress(0)
        status = st.empty()

        unique_frames = []
        previous_unique_frame = None
        images = []

        cols = [st.columns(3) for _ in range(len(frame_files) // 3 + 1)]
        col_index = 0
        row_index = 0

        for frame_index, frame_file in enumerate(frame_files):
            percentage_done = (frame_index / len(frame_files)) * 100
            progress_bar.progress(int(percentage_done))

            # Extract frame number from filename (frame_00123.jpg -> 123)
            frame_number_match = re.search(r"frame_(\d+)", frame_file)
            if not frame_number_match:
                continue
            frame_number = int(frame_number_match.group(1))

            # Load the frame
            frame_path = os.path.join(temp_all_folder, frame_file)
            frame = cv2.imread(frame_path)
            if frame is None:
                continue

            # Check if the frame is unique compared to the last unique frame
            if is_frame_unique_pixel(frame, previous_unique_frame, frame_index):
                print(f"Frame {frame_index} is unique!")
                previous_unique_frame = frame.copy()

                # Save to output directory
                output_frame_path = f"./output/frame_{frame_number}.jpg"
                cv2.imwrite(output_frame_path, frame)
                images.append(output_frame_path)
                status.text(f"Found unique frame {frame_index}")

                # Display in Streamlit
                if row_index < len(cols) and col_index < 3:
                    cols[row_index][col_index].image(output_frame_path, width=300)
                    col_index += 1
                    if col_index == 3:
                        col_index = 0
                        row_index += 1

        # Clean up temporary files
        os.unlink(temp_file_name)
        for file in os.listdir(temp_all_folder):
            os.unlink(os.path.join(temp_all_folder, file))
        os.rmdir(temp_all_folder)
        os.rmdir(temp_unique_folder)

        return images

    except Exception as e:
        print(f"Error during frame extraction: {e}")
        return []


def create_doc_with_images(image_paths):
    # Create a new Word document
    doc = Document()

    # Set the margins to 0.5 inches
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Loop over the image paths
    for i, image_path in enumerate(image_paths):
        # Add a new table for each image
        table = doc.add_table(rows=1, cols=2)

        for cell in table.columns[0].cells:
            cell.width = Inches(2.0)
        # Add the image to the first cell
        run = table.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(image_path, width=Inches(2.0))

        # Add text to the second cell
        table.cell(0, 1).text = "."

    # Save the document
    doc.save("./output/output.docx")


def process_file(uploaded_file):
    """Simulate file processing with progress."""

    placeholder_download = st.empty()
    placeholder_status = st.empty()
    images = extract_unique_frames(uploaded_file)
    print(images)
    placeholder_status.text(f"Found {len(images)} frames")
    st.text("Creating Word Doc")
    create_doc_with_images(images)

    st.success("Processing complete!")
    with open("./output/output.docx", "rb") as f:
        data = f.read()

    placeholder_download.download_button(
        label="Download Word Doc",
        data=data,
        file_name="output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def process_url(url):
    """Process URL to extract images and create document."""

    placeholder_download = st.empty()
    placeholder_status = st.empty()
    images = process_images_from_url(url)
    print(images)
    placeholder_status.text(f"Found {len(images)} unique images")
    st.text("Creating Word Doc")
    create_doc_with_images(images)

    st.success("Processing complete!")
    with open("./output/output.docx", "rb") as f:
        data = f.read()

    placeholder_download.download_button(
        label="Download Word Doc",
        data=data,
        file_name="output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def extract_comic_chapters(url):
    """Extract all chapters from a comic series page"""
    print(f"Extracting comic chapters from URL: {url}")
    
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36',
            'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
            'accept-language': 'en-US,en;q=0.9',
        }
        
        with httpx.Client(timeout=30.0) as client:
            response = client.get(url, headers=headers)
            response.raise_for_status()
            
            html_content = response.text
            
            # Look for chapter IDs in the HTML
            chapter_ids = []
            
            # Pattern 1: Look for [number] patterns - these are the actual chapter IDs
            import re
            id_matches = re.findall(r'\[(\d+)\]', html_content)
            print(f"Found {len(id_matches)} bracket pattern matches")
            for match in id_matches:
                if match not in chapter_ids:
                    chapter_ids.append(match)
                    print(f"Found chapter ID: {match}")
            
            # Pattern 2: Look for 7+ digit numbers (common chapter ID format)
            potential_ids = re.findall(r'(\d{7,})', html_content)  # 7+ digit numbers
            print(f"Found {len(potential_ids)} potential 7+ digit IDs")
            for potential_id in potential_ids:
                if potential_id not in chapter_ids:
                    chapter_ids.append(potential_id)
                    print(f"Found potential chapter ID: {potential_id}")
            
            # Pattern 3: Look for specific chapter patterns in the HTML
            chapter_patterns = [
                r'chapter-(\d+)',
                r'ch-(\d+)',
                r'episode-(\d+)',
                r'ep-(\d+)',
                r'issue-(\d+)',
                r'vol-(\d+)',
                r'volume-(\d+)'
            ]
            
            for pattern in chapter_patterns:
                matches = re.findall(pattern, html_content, re.IGNORECASE)
                for match in matches:
                    if match not in chapter_ids:
                        chapter_ids.append(match)
                        print(f"Found chapter ID with pattern {pattern}: {match}")
            
            # Pattern 4: Look for JavaScript data that might contain chapter IDs
            script_patterns = [
                r'chapters?\s*[:=]\s*\[([^\]]+)\]',
                r'chapterIds?\s*[:=]\s*\[([^\]]+)\]',
                r'episodes?\s*[:=]\s*\[([^\]]+)\]'
            ]
            
            for pattern in script_patterns:
                matches = re.findall(pattern, html_content, re.IGNORECASE)
                for match in matches:
                    # Extract numbers from the matched string
                    numbers = re.findall(r'\d+', match)
                    for number in numbers:
                        if number not in chapter_ids and len(number) >= 4:  # At least 4 digits
                            chapter_ids.append(number)
                            print(f"Found chapter ID in JS data: {number}")
            
            # Remove duplicates and sort
            chapter_ids = list(set(chapter_ids))
            chapter_ids.sort(key=int)
            
            print(f"Total chapter IDs found: {len(chapter_ids)}")
            return chapter_ids
            
    except Exception as e:
        st.error(f"Failed to extract chapters: {e}")
        return []

def fetch_comic_chapter_images(chapter_id, url):
    """Fetch images for a specific comic chapter"""
    print(f"Fetching images for chapter ID: {chapter_id}")
    
    # Use the exact headers from the curl command
    headers = {
        'accept': 'text/x-component',
        'accept-language': 'en-US,en;q=0.9',
        'cache-control': 'no-cache',
        'content-type': 'text/plain;charset=UTF-8',
        'next-action': '7f5de9b5f0064b81b6fe299caf58f1f4b5897eed51',
        'next-router-state-tree': '%5B%22%22%2C%7B%22children%22%3A%5B%5B%22domain%22%2C%22freecomic.to%22%2C%22d%22%5D%2C%7B%22children%22%3A%5B%22(routes)%22%2C%7B%22children%22%3A%5B%22series%22%2C%7B%22children%22%3A%5B%5B%22category%22%2C%22comic%22%2C%22d%22%5D%2C%7B%22children%22%3A%5B%5B%22slug%22%2C%22star-wars-legacy-of-vader-2025.164413%22%2C%22d%22%5D%2C%7B%22children%22%3A%5B%22__PAGE__%22%2C%7B%7D%2C%22%2Fcomic%2Fstar-wars-legacy-of-vader-2025.164413%22%2C%22refresh%22%5D%7D%5D%7D%5D%7D%5D%7D%5D%7D%5D%7D%2Cnull%2Cnull%2Ctrue%5D',
        'origin': 'https://freecomic.to',
        'pragma': 'no-cache',
        'priority': 'u=1, i',
        'referer': 'https://freecomic.to/comic/star-wars-legacy-of-vader-2025.164413',
        'sec-ch-ua': '"Not)A;Brand";v="8", "Chromium";v="138", "Google Chrome";v="138"',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '"macOS"',
        'sec-fetch-dest': 'empty',
        'sec-fetch-mode': 'cors',
        'sec-fetch-site': 'same-origin',
        'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
    }
    
    data = f'[{chapter_id}]'
    
    with httpx.Client(timeout=30.0) as client:
        try:
            response = client.post(url, headers=headers, data=data)
            response.raise_for_status()
            
            # Parse the response
            response_text = response.text
            
            # Parse the JSON response to extract image URLs
            try:
                import json
                # The response format is: "0:{...}\n1:[{...}]"
                response_lines = response_text.strip().split('\n')
                if len(response_lines) >= 2:
                    image_data_line = response_lines[1]
                    if image_data_line.startswith('1:'):
                        # Extract the JSON array part
                        json_part = image_data_line[2:]  # Remove "1:" prefix
                        image_array = json.loads(json_part)
                        
                        # Extract src URLs from each image object
                        image_urls = []
                        for image_obj in image_array:
                            if 'src' in image_obj:
                                image_urls.append(image_obj['src'])
                        
                        print(f"Found {len(image_urls)} images for chapter {chapter_id}")
                        return image_urls
                    else:
                        print("Unexpected response format")
                        
            except json.JSONDecodeError as e:
                print(f"Failed to parse JSON response: {e}")
                # Fallback to regex pattern matching
                import re
                imgor_pattern = r'https://imgor\.comick\.site/[^"\s]+'
                matches = re.findall(imgor_pattern, response_text)
                print(f"Found {len(matches)} imgor URLs using regex fallback")
                return matches
                
        except Exception as e:
            print(f"Error fetching chapter {chapter_id}: {e}")
            return []

def download_comic_image(url, output_path):
    """Download a single comic image and convert to JPEG"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
    }
    
    try:
        with httpx.Client(timeout=30.0) as client:
            response = client.get(url, headers=headers)
            response.raise_for_status()
            
            # Convert WebP to JPEG using Pillow
            from PIL import Image
            import io
            
            image = Image.open(io.BytesIO(response.content))
            
            # Convert to RGB if needed (WebP might be RGBA)
            if image.mode in ('RGBA', 'LA', 'P'):
                image = image.convert('RGB')
            
            # Save as JPEG
            jpeg_path = output_path + '.jpg'
            image.save(jpeg_path, 'JPEG', quality=95)
            
            return jpeg_path
            
    except Exception as e:
        print(f"Failed to download image {url}: {e}")
        return None

def process_comic_chapter(chapter_id, chapter_title, image_urls):
    """Process a comic chapter and create Word document using existing workflow"""
    print(f"Processing comic chapter: {chapter_title}")
    
    # Create temporary directory for downloaded images
    temp_dir = tempfile.mkdtemp(prefix="comic_images_")
    downloaded_images = []
    output_images = []
    
    # Create columns for image display
    num_cols = 3
    cols = st.columns(num_cols)
    
    st.text("Downloading and processing comic images")
    progress_bar = st.progress(0)
    status = st.empty()
    
    # Download and process images one by one, displaying as we go
    for i, url in enumerate(image_urls):
        percentage_done = (i / len(image_urls)) * 100
        progress_bar.progress(int(percentage_done))
        status.text(f"Downloading image {i+1}/{len(image_urls)}")
        
        print(f"Downloading image {i+1}/{len(image_urls)}: {url}")
        output_path = os.path.join(temp_dir, f"image_{i:04d}")
        downloaded_path = download_comic_image(url, output_path)
        
        if downloaded_path:
            downloaded_images.append(downloaded_path)
            
            # Process and display the image immediately
            frame = cv2.imread(downloaded_path)
            if frame is not None:
                # Save to output directory
                output_img_path = f"./output/comic_image_{i:04d}.jpg"
                cv2.imwrite(output_img_path, frame)
                output_images.append(output_img_path)
                
                # Display in Streamlit immediately
                col_index = i % num_cols
                with cols[col_index]:
                    st.image(output_img_path, width=300, caption=f"Image {i + 1}")
    
    if not output_images:
        print(f"No images downloaded for {chapter_title}")
        return None
    
    # Clean up temporary directory
    if temp_dir:
        for file in os.listdir(temp_dir):
            try:
                os.unlink(os.path.join(temp_dir, file))
            except:
                pass
        try:
            os.rmdir(temp_dir)
        except:
            pass
    
    return output_images

def process_chapter_url(chapter_url, chapter_title=None):
    """Process a specific chapter URL to extract images and create document."""

    placeholder_download = st.empty()
    placeholder_status = st.empty()
    
    if chapter_title:
        st.info(f"Processing chapter: {chapter_title}")
        st.warning("Note: This is a JavaScript-based chapter. The app will extract all images from the current page.")
    
    images = process_images_from_url(chapter_url)
    print(images)
    placeholder_status.text(f"Found {len(images)} unique images")
    st.text("Creating Word Doc")
    create_doc_with_images(images)

    st.success("Processing complete!")
    with open("./output/output.docx", "rb") as f:
        data = f.read()

    placeholder_download.download_button(
        label="Download Word Doc",
        data=data,
        file_name="output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


st.title("Video to Word")
st.subheader(
    "Upload a video or provide a URL to extract unique frames/images and save them to a Word document."
)

# Create tabs for different input methods
tab1, tab2 = st.tabs(["Comic Processing", "Upload Video"])

with tab1:
    st.write("### Comic Chapter Processing")
    st.write("Enter a comic series URL to extract chapters and download images.")
    
    comic_url = st.text_input("Enter comic series URL", placeholder="https://freecomic.to/comic/star-wars-legacy-of-vader-2025.164413")
    
    # Auto-extract chapters when URL is entered and user hits enter
    if comic_url:
        # Check if URL has changed or chapters haven't been extracted yet
        if ("comic_chapters" not in st.session_state or 
            st.session_state.comic_url != comic_url or 
            not st.session_state.comic_chapters):
            
            with st.spinner("Extracting comic chapters..."):
                chapter_ids = extract_comic_chapters(comic_url)
                st.session_state.comic_chapters = chapter_ids
                st.session_state.comic_url = comic_url
                st.session_state.comic_processed = False
            
            if not chapter_ids:
                st.warning("No comic chapters found. Please check the URL.")
            else:
                st.success(f"Found {len(chapter_ids)} comic chapters!")
        
        # Display comic chapters if available
        if "comic_chapters" in st.session_state and st.session_state.comic_chapters:
            st.write("### Available Comic Chapters:")
            for i, chapter_id in enumerate(st.session_state.comic_chapters):
                if st.button(f"Chapter {chapter_id}", key=f"comic_chapter_{i}"):
                    st.session_state.selected_comic_chapter = chapter_id
                    st.session_state.comic_processed = False
        
        # Process selected comic chapter
        if "selected_comic_chapter" in st.session_state:
            st.write(f"**Selected Comic Chapter:** {st.session_state.selected_comic_chapter}")
            
            # Auto-start processing when chapter is selected
            if "comic_processed" not in st.session_state or not st.session_state.comic_processed:
                with st.spinner("Fetching chapter images..."):
                    image_urls = fetch_comic_chapter_images(st.session_state.selected_comic_chapter, comic_url)
                    
                    if image_urls:
                        st.success(f"Found {len(image_urls)} images for chapter {st.session_state.selected_comic_chapter}")
                        
                        with st.spinner("Processing images..."):
                            chapter_title = f"Chapter_{st.session_state.selected_comic_chapter}"
                            output_images = process_comic_chapter(st.session_state.selected_comic_chapter, chapter_title, image_urls)
                            
                            if output_images:
                                st.text("Creating Word Doc")
                                create_doc_with_images(output_images)
                                
                                st.success("Processing complete!")
                                with open("./output/output.docx", "rb") as f:
                                    data = f.read()
                                
                                st.download_button(
                                    label="Download Word Document",
                                    data=data,
                                    file_name=f"chapter_{st.session_state.selected_comic_chapter}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                )
                            else:
                                st.error("Failed to process images.")
                    else:
                        st.error("No images found for this chapter.")
                
                st.session_state.comic_processed = True

with tab2:
    uploaded_file = st.file_uploader("Choose a video", type=['mp4', 'avi', 'mov', 'mkv'])

    if uploaded_file is not None:
        if "processed_video" not in st.session_state or not st.session_state.processed_video:
            process_file(uploaded_file)
            st.session_state.processed_video = True
