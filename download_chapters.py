import json
import httpx
import os
import tempfile
from docx import Document
from docx.shared import Inches
from urllib.parse import urlparse
from PIL import Image
import io

def download_image(url, output_path):
    """Download a single image and convert to JPEG if needed"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
    }
    
    try:
        with httpx.Client(timeout=30.0) as client:
            response = client.get(url, headers=headers)
            response.raise_for_status()
            
            # Convert WebP to JPEG using Pillow
            image = Image.open(io.BytesIO(response.content))
            
            # Convert to RGB if needed (WebP might be RGBA)
            if image.mode in ('RGBA', 'LA', 'P'):
                image = image.convert('RGB')
            
            # Save as JPEG
            jpeg_path = output_path + '.jpg'
            image.save(jpeg_path, 'JPEG', quality=95)
            
            return jpeg_path
            
    except Exception as e:
        print(f"Failed to download image {url}: {e}")
        return None

def create_word_document(image_paths, output_filename):
    """Create a Word document with all images"""
    doc = Document()
    
    # Set the margins to 0.5 inches
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Loop over the image paths
    for i, image_path in enumerate(image_paths):
        if not image_path or not os.path.exists(image_path):
            continue
            
        # Add a new table for each image
        table = doc.add_table(rows=1, cols=2)
        
        for cell in table.columns[0].cells:
            cell.width = Inches(2.0)
        
        # Add the image to the first cell
        run = table.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(image_path, width=Inches(2.0))
        
        # Add text to the second cell
        table.cell(0, 1).text = "."
    
    # Save the document
    doc.save(output_filename)
    print(f"Created Word document: {output_filename}")

def process_chapters():
    """Process all chapters from the JSON file"""
    print("Loading chapters from all_chapters.json...")
    
    try:
        with open("all_chapters.json", "r") as f:
            chapters = json.load(f)
    except FileNotFoundError:
        print("all_chapters.json not found. Run simple_extractor.py first.")
        return
    
    print(f"Found {len(chapters)} chapters to process")
    
    # Create output directory
    if not os.path.exists("output"):
        os.makedirs("output")
    
    for chapter in chapters:
        chapter_id = chapter['id']
        title = chapter['title']
        image_urls = chapter['image_urls']
        
        print(f"\nProcessing {title} ({len(image_urls)} images)...")
        
        # Create chapter directory
        chapter_dir = f"output/chapter_{chapter_id}"
        if not os.path.exists(chapter_dir):
            os.makedirs(chapter_dir)
        
        downloaded_images = []
        
        # Download all images for this chapter
        for i, url in enumerate(image_urls):
            print(f"Downloading image {i+1}/{len(image_urls)}: {url}")
            output_path = os.path.join(chapter_dir, f"image_{i:04d}")
            downloaded_path = download_image(url, output_path)
            
            if downloaded_path:
                downloaded_images.append(downloaded_path)
        
        # Create Word document for this chapter
        if downloaded_images:
            word_filename = os.path.join(chapter_dir, f"{title.replace(' ', '_')}.docx")
            create_word_document(downloaded_images, word_filename)
            print(f"Created Word document for {title}: {word_filename}")
        else:
            print(f"No images downloaded for {title}")
    
    print(f"\nProcessing complete! Check the 'output' directory for all chapter documents.")

if __name__ == "__main__":
    process_chapters() 